import streamlit as st
import pandas as pd
import time
import json
from PyPDF2 import PdfReader
from docx import Document as DocxReader
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from webdriver_manager.chrome import ChromeDriverManager
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import OpenAIEmbeddings
from langchain_community.chat_models import ChatOpenAI
from langchain.chains.question_answering import load_qa_chain
import plotly.graph_objects as go

# ---------------------- Dropdown of Countries ----------------------
countries = sorted([
    "United States", "Canada", "United Kingdom", "India", "Germany", "France",
    "Australia", "Singapore", "Netherlands", "Japan", "South Korea", "Brazil", "Mexico",
    "Spain", "Italy", "Sweden", "Norway", "Denmark", "Finland", "Switzerland", "Belgium",
    "Ireland", "United Arab Emirates", "South Africa", "New Zealand", "Remote"
])

# ---------------------- Resume Analyzer ----------------------
def resume_to_chunks(uploaded_file):
    text = ""
    if uploaded_file.name.endswith(".pdf"):
        reader = PdfReader(uploaded_file)
        text = "".join(page.extract_text() for page in reader.pages if page.extract_text())
    elif uploaded_file.name.endswith(".docx"):
        doc = DocxReader(uploaded_file)
        text = "\n".join([para.text for para in doc.paragraphs])
    else:
        st.error("Unsupported file format. Please upload PDF or DOCX.")
        return [], ""
    splitter = RecursiveCharacterTextSplitter(chunk_size=700, chunk_overlap=200)
    return splitter.split_text(text), text

from langchain.schema import Document

def run_openai_query(api_key, content, prompt):
    llm = ChatOpenAI(model="gpt-3.5-turbo", openai_api_key=api_key)
    chain = load_qa_chain(llm=llm, chain_type="stuff")

    # Handle string or list input
    if isinstance(content, str):
        input_documents = [Document(page_content=content)]
    else:
        input_documents = [Document(page_content=chunk) for chunk in content[:5]]

    return chain.run(input_documents=input_documents, question=prompt)


def get_strength_score(api_key, chunks):
    rubric_prompt = (
        "You are a resume reviewer. Please rate the resume on a scale from 0 to 100 using the rubric below:\n\n"
        "• 90-100: Outstanding resume — excellent technical skills, clear formatting, metrics-backed experience, highly relevant to target jobs.\n"
        "• 75-89: Strong resume — good clarity, relevant skills and roles, some quantification, minor improvements needed.\n"
        "• 60-74: Average — acceptable formatting and content but lacks metrics, specificity, or strong action verbs.\n"
        "• 40-59: Weak — lacks structure, missing key sections, vague or generic experience.\n"
        "• 0-39: Very poor — not a professional resume.\n\n"
        "Based on this rubric, provide only a number between 0 and 100. Do not explain or include any other text."
    )
    try:
        score_text = run_openai_query(api_key, chunks[:5], rubric_prompt)
        score = int("".join(filter(str.isdigit, score_text.strip())))
        return min(max(score, 0), 100)
    except Exception:
        return 70


def display_score_gauge(score, summary_text, api_key):
    if score >= 90:
        label = "Excellent"
        prompt = "In 2-3 sentences, explain why this resume is excellent without repeating the summary. Focus on technical skills, education, and experience."
    elif score >= 75:
        label = "Great"
        prompt = "In 2-3 sentences, explain why this resume is great. Focus on technical skills, education, and core strengths only."
    elif score >= 60:
        label = "Good"
        prompt = "Briefly explain why this resume is decent but can be improved, in a positive tone. Focus on education and technical skills."
    else:
        label = "Needs Improvement"
        prompt = "Write 2-3 sentences on why this resume needs improvement, without repeating the summary."

    try:
        response = run_openai_query(api_key, [summary_text], prompt)
    except:
        response = "This resume can be improved by aligning more closely with job-specific skills and metrics."

    # Create donut chart with dark theme
    fig = go.Figure(go.Pie(
        values=[score, 100 - score],
        hole=0.82,
        marker_colors=["#00D1C1", "#333333"],  # Turquoise for score, dark gray for remaining
        textinfo='none',
        hoverinfo='skip',
        showlegend=False
    ))

    fig.add_annotation(dict(
        text=f"<b>{score}</b>",
        x=0.5, y=0.5,
        showarrow=False,
        font_size=64,
        font=dict(color="#FFFFFF")  # White font color for score
    ))

    fig.update_traces(marker=dict(line=dict(color='#121212', width=2)))  # Dark border
    fig.update_layout(
        margin=dict(t=10, b=10, l=0, r=0), 
        width=220, 
        height=220,
        paper_bgcolor='rgba(0,0,0,0)',  # Transparent background
        plot_bgcolor='rgba(0,0,0,0)'    # Transparent plot area
    )

    col1, col2 = st.columns([0.6, 1.4])
    with col1:
        st.plotly_chart(fig, use_container_width=False)

    with col2:
        st.markdown(f"<h2 style='margin-top:5px; color:white;'>{label}</h2>", unsafe_allow_html=True)
        st.markdown(f"<div style='font-size:16px; color:#e0e0e0; line-height:1.5;'>{response}</div>", unsafe_allow_html=True)

def get_resume_fixes(api_key, chunks):
    prompt = (
        "Analyze the resume and return a JSON list of exactly these 4 standard improvement categories: "
        "Weak Verbs, Buzzwords, Filler Words, Consistency. "
        "For each, provide: issue name, severity score (1-10), examples with suggestions. "
        "For Weak Verbs: identify generic verbs that could be replaced with stronger action verbs. "
        "For Buzzwords: identify overused industry jargon or trendy terms. "
        "For Filler Words: identify unnecessary words that add no value. "
        "For Consistency: identify any inconsistencies in formatting, tense, or style. "
        "Format: [{\"issue\": \"Weak Verbs\", \"score\": 6, \"details\": [{\"word\": \"helped with\", \"suggestion\": \"Replace with 'spearheaded' or 'led'\"}]}]"
    )
    try:
        raw = run_openai_query(api_key, chunks, prompt)
        return json.loads(raw)
    except Exception as e:
        # Fallback data if parsing fails
        return [
            {
                "issue": "Weak Verbs", 
                "score": 7, 
                "details": [
                    {"word": "worked on", "suggestion": "Replace with 'developed', 'implemented', or 'executed'"},
                    {"word": "helped with", "suggestion": "Replace with 'led', 'coordinated', or 'orchestrated'"}
                ]
            },
            {
                "issue": "Buzzwords", 
                "score": 5, 
                "details": [
                    {"word": "synergy", "suggestion": "Replace with specific collaborative achievements"},
                    {"word": "results-driven", "suggestion": "Include actual metrics and outcomes instead"}
                ]
            },
            {
                "issue": "Filler Words", 
                "score": 4, 
                "details": [
                    {"word": "very", "suggestion": "Remove or replace with specific descriptors"},
                    {"word": "in order to", "suggestion": "Replace with 'to' for conciseness"}
                ]
            },
            {
                "issue": "Consistency", 
                "score": 6, 
                "details": [
                    {"word": "Mixed tenses", "suggestion": "Use past tense for previous roles and present for current roles"},
                    {"word": "Inconsistent formatting", "suggestion": "Standardize bullet points, dates, and section headers"}
                ]
            }
        ]

def display_top_fixes(fixes):
    st.markdown("## Top Fixes")
    
    # Apply custom CSS to match dark theme
    st.markdown("""
    <style>
    .stExpander {
        background-color: #1e1e1e !important;
        border: 1px solid #333 !important;
        border-radius: 8px !important;
        margin-bottom: 10px !important;
        color: white !important;
    }
    .stExpander > div[role="button"] {
        color: white !important;
        font-weight: 500 !important;
    }
    .stExpander > div[role="button"]:hover {
        background-color: #2d2d2d !important;
    }
    </style>
    """, unsafe_allow_html=True)
    
    for fix in fixes:
        score_color = "#e74c3c" if fix["score"] >= 7 else "#f39c12" if fix["score"] >= 4 else "#2ecc71"
        with st.expander(fix['issue']):
            # Severity indicator on the right side
            st.markdown(
                f"<div style='display:flex; justify-content:flex-end;'>"
                f"<span style='background-color:{score_color}; color:white; padding:4px 10px; "
                f"border-radius:20px; font-weight:bold;'>{fix['score']}</span></div>",
                unsafe_allow_html=True
            )
            
            if fix.get("details"):
                for detail in fix["details"]:
                    st.markdown(
                        f"<div style='margin-top:10px; margin-bottom:10px; padding:10px; background-color:#2d2d2d; border-radius:5px;'>"
                        f"<strong style='color:#ddd;'>{detail.get('word', '—')}</strong><br>"
                        f"<span style='color:#00D1C1; font-style:italic;'>{detail.get('suggestion', 'Consider revising')}</span>"
                        f"</div>",
                        unsafe_allow_html=True
                    )
            else:
                st.write("No specific examples detected.")

# ---------------------- LinkedIn Scraper ----------------------
def scrape_jobs(role, location, resume_summary, num_jobs=10):
    from selenium.webdriver.common.action_chains import ActionChains

    options = Options()
    options.add_argument("--start-maximized")
    options.add_experimental_option("detach", True)
    driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)

    st.session_state["driver"] = driver
    driver.get("https://www.linkedin.com/login")

    st.warning("🔐 Please log into LinkedIn manually in the browser that just opened.")
    st.info("⏳ Waiting for you to log in... Scraping will start once login is successful.")

    while "login" in driver.current_url:
        time.sleep(2)

    query = f"https://www.linkedin.com/jobs/search/?keywords={role}&location={location}&sortBy=DD"
    driver.get(query)
    driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
    time.sleep(5)

    try:
        WebDriverWait(driver, 60).until(
            EC.presence_of_all_elements_located((By.CLASS_NAME, "job-card-container"))
        )
        time.sleep(3)

        job_cards = driver.find_elements(By.CLASS_NAME, "job-card-container")
        jobs_to_scrape = min(num_jobs, len(job_cards))
        jobs = []

        for i, card in enumerate(job_cards[:jobs_to_scrape]):
            try:
                driver.execute_script("arguments[0].scrollIntoView();", card)
                time.sleep(1)
                ActionChains(driver).move_to_element(card).click().perform()
                time.sleep(2)

                # Extract URL from job card element
                try:
                    job_url = card.find_element(By.CSS_SELECTOR, "a").get_attribute("href")
                except:
                    job_url = "URL not found"

                if "linkedin.com/jobs/view/" in job_url:
                    jobs.append({"Job URL": job_url})
                else:
                    st.warning(f"Job {i+1} does not have a valid job view URL.")

            except Exception as e:
                st.error(f"❌ Error scraping job {i+1}: {e}")
                with open(f"job_error_{i+1}.html", "w", encoding="utf-8") as f:
                    f.write(driver.page_source)

        df = pd.DataFrame(jobs)
        if not df.empty:
            st.success(f"✅ Scraped {len(df)} job URLs.")
            st.dataframe(df)
            st.download_button("⬇️ Download Job URLs as CSV", df.to_csv(index=False), file_name="linkedin_job_urls.csv", mime="text/csv")
        else:
            st.warning("❗No job data collected. Try again or verify LinkedIn loaded correctly.")

        return df

    finally:
        if "driver" in st.session_state:
            del st.session_state["driver"]
            driver.quit()

# ---------------------- Streamlit UI ----------------------
st.set_page_config(page_title="Resume & LinkedIn Analyzer", layout="wide")

# Apply custom dark theme styling
st.markdown("""
<style>
    body {
        background-color: #121212;
        color: #f0f0f0;
    }
    .stApp {
        background-color: #121212;
    }
    h1, h2, h3, h4, h5, h6 {
        color: white !important;
    }
    .stTabs [data-baseweb="tab-list"] {
        background-color: #1e1e1e;
        border-radius: 8px;
    }
    .stTabs [data-baseweb="tab"] {
        color: #f0f0f0;
    }
    .stTabs [aria-selected="true"] {
        background-color: #2d2d2d !important;
        color: #00D1C1 !important;
    }
    .stButton>button {
        background-color: #00D1C1;
        color: #121212;
        border: none;
        border-radius: 4px;
        padding: 0.5rem 1rem;
        font-weight: bold;
    }
    .stButton>button:hover {
        background-color: #00a399;
    }
    .stTextInput>div>div {
        background-color: #2d2d2d;
        color: white;
    }
    .stDataFrame {
        background-color: #2d2d2d;
    }
    .stDownloadButton>button {
        background-color: #2d2d2d;
        color: #00D1C1;
        border: 1px solid #00D1C1;
    }
</style>
""", unsafe_allow_html=True)

st.title("📄 AI Resume Analyzer + 💼 LinkedIn Job Scraper")

tabs = st.tabs(["Resume Analyzer", "LinkedIn Jobs"])

# Initialize a dictionary to store resume versions
resume_versions = {}

# Resume Analyzer Tab
with tabs[0]:
    st.subheader("Resume Analysis")
    openai_api_key = "Your API Key"
    uploaded_files = st.file_uploader("📤 Upload your Resume (PDF or DOCX)", type=["pdf", "docx"], accept_multiple_files=True)

    if uploaded_files and openai_api_key:
        for file in uploaded_files:
            chunks, full_text = resume_to_chunks(file)
            with st.spinner(f"Analyzing {file.name}..."):
                summary = run_openai_query(openai_api_key, chunks, "Summarize this resume.")
                strengths = run_openai_query(openai_api_key, chunks, "What are the strengths in this resume?")
                score = get_strength_score(openai_api_key, chunks)
                weaknesses = run_openai_query(openai_api_key, chunks, "What are the weaknesses in this resume?")
                roles = run_openai_query(openai_api_key, chunks, "Based on this resume, what job roles are suitable?")
                
                resume_versions[file.name] = {
                    "score": score, 
                    "summary": summary, 
                    "chunks": chunks, 
                    "strengths": strengths,
                    "weaknesses": weaknesses,
                    "roles": roles
                }
                
        # Set the first resume as default for LinkedIn job matching
        if uploaded_files:
            first_file = uploaded_files[0]
            st.session_state["resume_summary"] = resume_versions[first_file.name]["summary"]
        
                    # Display single resume analysis
        if len(uploaded_files) == 1:
            file = uploaded_files[0]
            
            # Create a container with custom styling
            with st.container():
                st.markdown(
                    f"""
                    <div style="background-color:#1e1e1e; padding:15px; border-radius:10px; margin-bottom:20px;">
                        <div style="display:flex; align-items:center;">
                            <div style="margin-right:10px;">
                                <svg width="24" height="24" viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
                                    <path d="M14 2H6C4.9 2 4 2.9 4 4V20C4 21.1 4.9 22 6 22H18C19.1 22 20 21.1 20 20V8L14 2Z" stroke="#00D1C1" stroke-width="2" stroke-miterlimit="10" stroke-linecap="round" stroke-linejoin="round"/>
                                    <path d="M14 2V8H20" stroke="#00D1C1" stroke-width="2" stroke-miterlimit="10" stroke-linecap="round" stroke-linejoin="round"/>
                                </svg>
                            </div>
                            <h3 style="margin:0; color:white;">{file.name}</h3>
                        </div>
                    </div>
                    """, 
                    unsafe_allow_html=True
                )
            
            # Display score gauge
            display_score_gauge(resume_versions[file.name]["score"], resume_versions[file.name]["summary"], openai_api_key)
            
            # Display top fixes
            display_top_fixes(get_resume_fixes(openai_api_key, resume_versions[file.name]["chunks"]))
            
            # Create collapsible sections for detailed analysis
            with st.expander("📝 Summary", expanded=False):
                st.markdown(f'<div style="background-color:#2d2d2d; padding:15px; border-radius:8px;">{resume_versions[file.name]["summary"]}</div>', unsafe_allow_html=True)
                
            with st.expander("💪 Strengths", expanded=False):
                st.markdown(f'<div style="background-color:#2d2d2d; padding:15px; border-radius:8px;">{resume_versions[file.name]["strengths"]}</div>', unsafe_allow_html=True)
                
            with st.expander("⚠️ Weaknesses", expanded=False):
                st.markdown(f'<div style="background-color:#2d2d2d; padding:15px; border-radius:8px;">{resume_versions[file.name]["weaknesses"]}</div>', unsafe_allow_html=True)
                
            with st.expander("🎯 Job Role Suggestions", expanded=False):
                st.markdown(f'<div style="background-color:#2d2d2d; padding:15px; border-radius:8px;">{resume_versions[file.name]["roles"]}</div>', unsafe_allow_html=True)
            
        # Display comparison for multiple resumes
        elif len(uploaded_files) > 1:
            st.subheader("Resume Score Comparison")
            scores = {name: resume_versions[name]["score"] for name in resume_versions}
            st.bar_chart(scores)
            
            # Allow user to select which resume to view in detail
            selected_resume = st.selectbox("Select resume to view details", list(resume_versions.keys()))
            
            if selected_resume:
                display_score_gauge(resume_versions[selected_resume]["score"], resume_versions[selected_resume]["summary"], openai_api_key)
                display_top_fixes(get_resume_fixes(openai_api_key, resume_versions[selected_resume]["chunks"]))
                
                st.write("### 📝 Summary")
                st.write(resume_versions[selected_resume]["summary"])
                st.write("### 💪 Strengths")
                st.write(resume_versions[selected_resume]["strengths"])
                st.write("### ⚠️ Weaknesses")
                st.write(resume_versions[selected_resume]["weaknesses"])
                st.write("### 🎯 Job Role Suggestions")
                st.write(resume_versions[selected_resume]["roles"])

# LinkedIn Jobs Tab
with tabs[1]:
    st.subheader("Search LinkedIn Jobs")
    job_title = st.text_input("Enter Job Title", "Data Scientist")
    job_location = st.selectbox("Select Job Location", countries)

    if st.button("Search Jobs"):
        if "resume_summary" not in st.session_state:
            st.error("Please analyze a resume first to use it for matching.")
        else:
            with st.spinner("Opening LinkedIn... Please log in manually."):
                df = scrape_jobs(job_title, job_location, st.session_state["resume_summary"], 25)
                if df.empty:
                    st.warning("No matching jobs found. Try changing the title or location.")
                else:
                    st.session_state["jobs_df"] = df
